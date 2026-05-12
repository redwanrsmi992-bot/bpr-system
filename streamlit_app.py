import streamlit as st
import pandas as pd
import plotly.graph_objs as go
from flask import Flask
from models import db, Employee, Process, Step
import io

# ---- حماية بكلمة مرور ----
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("نظام إعادة هندسة العمليات")
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("### دخول اعضاء الفريق")
        password = st.text_input("ادخل كلمة المرور", type="password")
        if st.button("دخول", use_container_width=True):
            if password == "BPR2026":
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("كلمة المرور خاطئة")
        st.markdown("---")
        st.caption("للاستفسار عن كلمة المرور، تواصل مع مسؤول النظام")
    st.stop()

# ---- التطبيق الرئيسي ----
st.set_page_config(page_title="نظام إعادة هندسة العمليات", layout="wide")
# ============ تحسين الواجهة ============
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Tajawal', sans-serif;
    }
    
    .main {
        background-color: #f8fafc;
    }
    
    .stButton>button {
        border-radius: 8px;
        background-color: #2563eb;
        color: white;
        font-weight: bold;
        border: none;
        padding: 8px 16px;
        transition: all 0.2s ease;
    }
    
    .stButton>button:hover {
        background-color: #1d4ed8;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transform: translateY(-1px);
    }
    
    .stMetric {
        background-color: white;
        padding: 16px;
        border-radius: 12px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        text-align: center;
    }
    
    .stExpander {
        background-color: white;
        border-radius: 12px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        border: 1px solid #e2e8f0;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: #f1f5f9;
        border-radius: 8px;
        padding: 4px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 6px;
        padding: 8px 16px;
        font-weight: bold;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #2563eb !important;
        color: white !important;
    }
    
    h1, h2, h3 {
        color: #1e293b;
        font-weight: bold;
    }
    
    .stAlert {
        border-radius: 8px;
    }
    
    .stDataFrame {
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    ::-webkit-scrollbar {
        width: 8px;
    }
    ::-webkit-scrollbar-track {
        background: #f1f5f9;
    }
    ::-webkit-scrollbar-thumb {
        background: #94a3b8;
        border-radius: 4px;
    }
</style>
""", unsafe_allow_html=True)
# =====================================
st.title(" نظام إعادة هندسة العمليات ")

if st.sidebar.button("تسجيل الخروج"):
    st.session_state.authenticated = False
    st.rerun()

# ---- إعداد قاعدة البيانات ----
app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///bpr_system.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db.init_app(app)

with app.app_context():
    db.create_all()

# ---- دوال مساعدة ----
def get_processes():
    with app.app_context():
        processes = Process.query.all()
        for p in processes:
            _ = p.steps
        return processes

def get_process_by_id(pid):
    with app.app_context():
        p = Process.query.get(pid)
        if p:
            _ = p.steps
        return p

def get_steps(pid):
    with app.app_context():
        return Step.query.filter_by(process_id=pid).order_by(Step.step_order).all()

def add_process_to_db(name, category, freq, status):
    with app.app_context():
        p = Process(name=name, category=category, annual_frequency=freq, status=status)
        db.session.add(p)
        db.session.commit()

def update_process_in_db(pid, name, category, freq, status):
    with app.app_context():
        p = Process.query.get(pid)
        if p:
            p.name = name
            p.category = category
            p.annual_frequency = freq
            p.status = status
            db.session.commit()
            return True
        return False

def delete_process_from_db(pid):
    with app.app_context():
        p = Process.query.get(pid)
        if p:
            db.session.delete(p)
            db.session.commit()
            return True
        return False

def add_employee_to_db(title, cost):
    with app.app_context():
        e = Employee(title=title, monthly_cost=cost)
        db.session.add(e)
        db.session.commit()

def update_employee_in_db(eid, title, cost):
    with app.app_context():
        e = Employee.query.get(eid)
        if e:
            e.title = title
            e.monthly_cost = cost
            db.session.commit()
            return True
        return False

def delete_employee_from_db(eid):
    with app.app_context():
        e = Employee.query.get(eid)
        if e:
            db.session.delete(e)
            db.session.commit()
            return True
        return False

def get_employees():
    with app.app_context():
        return Employee.query.all()

def add_step_to_db(pid, eid, order, name, pt, wt, stype, sys_used, waste):
    with app.app_context():
        s = Step(process_id=pid, employee_id=eid, step_order=order,
                 step_name=name, processing_time_minutes=pt,
                 wait_time_minutes=wt, step_type=stype,
                 system_used=sys_used, waste_category=waste)
        db.session.add(s)
        db.session.commit()

# ---- القائمة الجانبية ----
menu = st.sidebar.radio("القائمة", [
    "الرئيسية",
    "اضافة موظف",
    "اضافة عملية",
    "اضافة خطوات",
    "لوحة التحكم",
    "رحلة متلقي الخدمة",
    "مصفوفة الاثر والتاثير",
    "رفع ملف عمليات",
    "🎯 تحليل باريتو (80/20)",
    "📋 SIPOC",
    "📊 RACI",
    "🗺️ الخريطة الحرارية",
    "🚀 توصيات التحسين",
    "📊 مخطط BPMN",
    "📄 تقرير العملية",
    "📄 تقرير PDF",
    "📄 رفع نموذج Word",
    "دليل الاستخدام"
])
# ================== الصفحة الرئيسية ==================
if menu == "الرئيسية":
    st.subheader("قائمة العمليات")
    
    processes = get_processes()
    if processes:
        for p in processes:
            with app.app_context():
                eff = Process.query.get(p.id).flow_efficiency
                cost = Process.query.get(p.id).annual_cost
            with st.expander(f"{p.id} - {p.name} ({p.category})"):
                col1, col2, col3 = st.columns(3)
                col1.metric("كفاءة التدفق", f"{eff:.2f}%")
                col2.metric("التكلفة السنوية", f"{cost:,.2f} د.ا")
                col3.metric("التكرار السنوي", p.annual_frequency)
    else:
        st.info("لا توجد عمليات بعد.")

    # ---- استيراد وتصدير ----
    with st.expander("💾 حفظ واستعادة البيانات", expanded=False):
        st.markdown("لتجنب فقدان البيانات، حمّل نسخة احتياطية قبل إعادة النشر.")
        col_backup, col_restore = st.columns(2)
        
        with col_backup:
            if st.button("📥 تحميل نسخة احتياطية (Excel)"):
                with app.app_context():
                    emp_data = [{"id": e.id, "title": e.title, "cost": e.monthly_cost} for e in Employee.query.all()]
                    proc_data = [{"id": p.id, "name": p.name, "category": p.category, "freq": p.annual_frequency, "status": p.status} for p in Process.query.all()]
                    step_data = [{"id": s.id, "process_id": s.process_id, "employee_id": s.employee_id, "order": s.step_order, "name": s.step_name, "pt": s.processing_time_minutes, "wt": s.wait_time_minutes, "type": s.step_type, "system": s.system_used, "waste": s.waste_category} for s in Step.query.all()]
                    
                    df_emp = pd.DataFrame(emp_data) if emp_data else pd.DataFrame()
                    df_proc = pd.DataFrame(proc_data) if proc_data else pd.DataFrame()
                    df_step = pd.DataFrame(step_data) if step_data else pd.DataFrame()
                    
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        if not df_emp.empty: df_emp.to_excel(writer, sheet_name='employees', index=False)
                        if not df_proc.empty: df_proc.to_excel(writer, sheet_name='processes', index=False)
                        if not df_step.empty: df_step.to_excel(writer, sheet_name='steps', index=False)
                    output.seek(0)
                    
                    st.download_button("⬇️ تحميل النسخة الاحتياطية", data=output, file_name="bpr_backup.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        with col_restore:
            uploaded_backup = st.file_uploader("📤 استعادة من نسخة احتياطية", type="xlsx", key="restore")
            if uploaded_backup is not None:
                if st.button("🔄 استعادة البيانات"):
                    try:
                        with app.app_context():
                            xl = pd.ExcelFile(uploaded_backup)
                            
                            if 'employees' in xl.sheet_names:
                                df_e = pd.read_excel(uploaded_backup, sheet_name='employees')
                                for _, row in df_e.iterrows():
                                    if not Employee.query.get(row['id']):
                                        emp = Employee(id=row['id'], title=row['title'], monthly_cost=row['cost'])
                                        db.session.add(emp)
                            
                            if 'processes' in xl.sheet_names:
                                df_p = pd.read_excel(uploaded_backup, sheet_name='processes')
                                for _, row in df_p.iterrows():
                                    if not Process.query.get(row['id']):
                                        proc = Process(id=row['id'], name=row['name'], category=row['category'], annual_frequency=row['freq'], status=row['status'])
                                        db.session.add(proc)
                            
                            if 'steps' in xl.sheet_names:
                                df_s = pd.read_excel(uploaded_backup, sheet_name='steps')
                                for _, row in df_s.iterrows():
                                    if not Step.query.get(row['id']):
                                        step = Step(id=row['id'], process_id=row['process_id'], employee_id=row['employee_id'], step_order=row['order'], step_name=row['name'], processing_time_minutes=row['pt'], wait_time_minutes=row['wt'], step_type=row['type'], system_used=row['system'], waste_category=row['waste'])
                                        db.session.add(step)
                            
                            db.session.commit()
                            st.success("تم استعادة جميع البيانات بنجاح!")
                            st.rerun()
                    except Exception as e:
                        st.error(f"فشل الاستيراد: {e}")


# ================== اضافة موظف ==================
elif menu == "اضافة موظف":
    st.subheader("ادارة الموظفين")
    
    with st.expander("اضافة موظف جديد", expanded=False):
        with st.form("add_emp"):
            title = st.text_input("المسمى الوظيفي")
            cost = st.number_input("الراتب الشهري (دينار)", min_value=100, value=500)
            if st.form_submit_button("حفظ"):
                if title:
                    add_employee_to_db(title, float(cost))
                    st.success(f"تمت اضافة الموظف: {title}")
                    st.rerun()
                else:
                    st.error("الرجاء ادخال المسمى الوظيفي")
    
    emps = get_employees()
    if emps:
        st.subheader("الموظفون الحاليون")
        for e in emps:
            with st.expander(f"{e.id} - {e.title} ({e.monthly_cost:.0f} د.ا)"):
                if st.button("حذف", key=f"del_emp_{e.id}"):
                    delete_employee_from_db(e.id)
                    st.success("تم الحذف")
                    st.rerun()

# ================== اضافة عملية ==================
elif menu == "اضافة عملية":
    st.subheader("اضافة عملية جديدة")
    with st.form("add_proc"):
        name = st.text_input("اسم العملية")
        category = st.selectbox("الفئة", ["استراتيجية", "انتصار_سريع", "روتينية", "للدراسة"])
        freq = st.number_input("التكرار السنوي", min_value=1, value=1)
        status = st.selectbox("الحالة", ["غير_مبدوء", "تحت_الدراسة", "مكتمل"])
        if st.form_submit_button("حفظ"):
            if name:
                add_process_to_db(name, category, freq, status)
                st.success(f"تمت اضافة العملية: {name}")
                st.rerun()
            else:
                st.error("الرجاء ادخال اسم العملية")

# ================== اضافة خطوات (آمن) ==================
elif menu == "اضافة خطوات":
    st.subheader("ادارة الخطوات")

    tab1, tab2 = st.tabs(["➕ اضافة خطوة", "📋 عرض وتعديل الخطوات"])

    # ---- تبويبة الاضافة ----
    with tab1:
        processes = get_processes()
        employees = get_employees()
        if not processes:
            st.warning("لا توجد عمليات")
        elif not employees:
            st.warning("لا توجد موظفين")
        else:
            pnames = [f"{p.id} - {p.name}" for p in processes]
            enames = [f"{e.id} - {e.title}" for e in employees]
            
            # النموذج الأساسي (Form)
            with st.form("add_step_form"):
                pid_sel = st.selectbox("اختر العملية", pnames)
                eid_sel = st.selectbox("اختر الموظف", enames)
                order = st.number_input("رقم الترتيب", min_value=1, value=1)
                sname = st.text_input("اسم الخطوة")
                pt = st.number_input("وقت المعالجة (دقيقة)", min_value=0.0, value=5.0)
                wt = st.number_input("وقت الانتظار (دقيقة)", min_value=0.0, value=0.0)
                stype = st.selectbox("نوع الخطوة", ["VA", "BNVA", "NVA"])
                system = st.selectbox("النظام المستخدم", ["Oracle", "GFMIS", "Outlook", "ورقي", "يدوي"])
                waste = st.text_input("فئة الهدر (ان وجدت)")
                
                # زر الحفظ (داخل النموذج)
                if st.form_submit_button("💾 حفظ الخطوة"):
                    if sname:
                        pid = int(pid_sel.split(" - ")[0])
                        eid = int(eid_sel.split(" - ")[0])
                        add_step_to_db(pid, eid, order, sname, pt, wt, stype, system, waste)
                        st.success("تمت اضافة الخطوة")
                        st.rerun()
                    else:
                        st.error("الرجاء ادخال اسم الخطوة")
            
            # المساعد (خارج النموذج
            with st.expander("💡 تحليل نوع الخطوة (مساعد)"):
                if sname:
                    st.markdown("""
                    **دليل سريع للتصنيف:**
                    - **VA (قيمة مضافة):** متلقي الخدمة يدفع مقابل هذه الخطوة (مثلاً: تسجيل طلب، إصدار أمر دفع).
                    - **BNVA (هدر ضروري):** إجراء رقابي أو قانوني (مثلاً: تدقيق الرصيد، مراجعة مدير).
                    - **NVA (هدر خالص):** انتظار، إعادة عمل، موافقات زائدة.
                    
                    **اسأل نفسك:** "لو كنت المواطن، هل سأدفع مقابل هذه الخطوة؟"
                    - إذا كان الجواب **نعم** → VA
                    - إذا كان الجواب **لا، لكنها إجبارية** → BNVA
                    - إذا كان الجواب **لا، ويمكن الاستغناء عنها** → NVA
                    """)

    # ---- تبويبة العرض والتعديل ----
    with tab2:
        processes = get_processes()
        if processes:
            pnames = [f"{p.id} - {p.name}" for p in processes]
            sel_process = st.selectbox("اختر العملية لعرض خطواتها", pnames, key="view_steps")
            pid = int(sel_process.split(" - ")[0])
            steps = get_steps(pid)
            
            if steps:
                st.markdown(f"**عدد الخطوات:** {len(steps)}")
                for s in steps:
                    with app.app_context():
                        emp = Employee.query.get(s.employee_id)
                        emp_title = emp.title if emp else "-"
                    with st.expander(f"{s.step_order}. {s.step_name} | {s.step_type} | {emp_title}"):
                        col_info, col_actions = st.columns([3, 1])
                        with col_info:
                            st.markdown(f"""
                            - **الموظف:** {emp_title}
                            - **وقت العمل:** {s.processing_time_minutes} دقيقة
                            - **وقت الانتظار:** {s.wait_time_minutes} دقيقة
                            - **النظام:** {s.system_used or '-'}
                            - **الهدر:** {s.waste_category or '-'}
                            """)
                        with col_actions:
                            # تعديل وحذف (خارج form)
                            if st.button("✏️ تعديل", key=f"edit_step_{s.id}"):
                                st.session_state[f"editing_step_{s.id}"] = True
                            if st.button("🗑️ حذف", key=f"del_step_{s.id}"):
                                with app.app_context():
                                    step = Step.query.get(s.id)
                                    if step:
                                        db.session.delete(step)
                                        db.session.commit()
                                st.success("تم الحذف")
                                st.rerun()
                        
                        if st.session_state.get(f"editing_step_{s.id}", False):
                            with st.form(f"edit_step_form_{s.id}"):
                                employees = get_employees()
                                enames = [f"{e.id} - {e.title}" for e in employees]
                                current_emp = f"{s.employee_id} - {emp_title}" if s.employee_id else enames[0]
                                new_eid_sel = st.selectbox("الموظف", enames, index=enames.index(current_emp) if current_emp in enames else 0)
                                new_order = st.number_input("رقم الترتيب", min_value=1, value=s.step_order)
                                new_name = st.text_input("اسم الخطوة", value=s.step_name)
                                new_pt = st.number_input("وقت المعالجة", min_value=0.0, value=s.processing_time_minutes)
                                new_wt = st.number_input("وقت الانتظار", min_value=0.0, value=s.wait_time_minutes)
                                new_type = st.selectbox("النوع", ["VA", "BNVA", "NVA"], index=["VA", "BNVA", "NVA"].index(s.step_type))
                                new_system = st.selectbox("النظام", ["Oracle", "GFMIS", "Outlook", "ورقي", "يدوي"], index=["Oracle", "GFMIS", "Outlook", "ورقي", "يدوي"].index(s.system_used) if s.system_used in ["Oracle", "GFMIS", "Outlook", "ورقي", "يدوي"] else 0)
                                new_waste = st.text_input("فئة الهدر", value=s.waste_category or "")
                                
                                col_save, col_cancel = st.columns(2)
                                with col_save:
                                    if st.form_submit_button("💾 حفظ"):
                                        with app.app_context():
                                            step = Step.query.get(s.id)
                                            if step:
                                                step.employee_id = int(new_eid_sel.split(" - ")[0])
                                                step.step_order = new_order
                                                step.step_name = new_name
                                                step.processing_time_minutes = new_pt
                                                step.wait_time_minutes = new_wt
                                                step.step_type = new_type
                                                step.system_used = new_system
                                                step.waste_category = new_waste
                                                db.session.commit()
                                        st.session_state[f"editing_step_{s.id}"] = False
                                        st.success("تم التعديل")
                                        st.rerun()
                                with col_cancel:
                                    if st.form_submit_button("❌ الغاء"):
                                        st.session_state[f"editing_step_{s.id}"] = False
                                        st.rerun()
            else:
                st.info("لا توجد خطوات لهذه العملية")
        else:
            st.info("لا توجد عمليات بعد")
# ================== لوحة القيادة (النهائية مع وقت الانتظار) ==================
elif menu == "لوحة التحكم":
    st.subheader("📊 لوحة القيادة (Executive Dashboard)")
    st.markdown("نظرة شاملة على أداء جميع العمليات في الدائرة.")

    all_processes = get_processes()
    
    if all_processes:
        # --- 1. بطاقات الملخص العام (Summary KPIs) ---
        total_processes = len(all_processes)
        total_waste_minutes = 0
        total_processing_minutes = 0
        total_annual_cost_comprehensive = 0
        
        for proc in all_processes:
            with app.app_context():
                p = Process.query.get(proc.id)
                wait = sum((s.wait_time_minutes or 0) for s in p.steps)
                proc_time = sum((s.processing_time_minutes or 0) for s in p.steps)
                total_waste_minutes += wait
                total_processing_minutes += proc_time
                total_annual_cost_comprehensive += (proc_time + wait) * 0.1 * p.annual_frequency

        avg_flow_eff = (total_processing_minutes / (total_processing_minutes + total_waste_minutes) * 100) if (total_processing_minutes + total_waste_minutes) > 0 else 0
        
        st.markdown("### 🎯 ملخص الأداء العام")
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("📋 إجمالي العمليات", total_processes)
        col2.metric("⚡ متوسط كفاءة التدفق", f"{avg_flow_eff:.1f}%")
        col3.metric("⏳ إجمالي وقت الهدر", f"{total_waste_minutes:,.0f} دقيقة")
        col4.metric("💸 إجمالي التكلفة الشاملة", f"{total_annual_cost_comprehensive:,.2f} د.أ")

        st.markdown("---")

        # --- 2. جدول ملخص العمليات مع وقت الانتظار ---
        st.subheader("📋 ملخص جميع العمليات")
        summary_data = []
        for proc in all_processes:
            with app.app_context():
                p = Process.query.get(proc.id)
                wait = sum((s.wait_time_minutes or 0) for s in p.steps)
                proc_time = sum((s.processing_time_minutes or 0) for s in p.steps)
                lead_time = proc_time + wait
                flow_eff = (proc_time / lead_time * 100) if lead_time > 0 else 100
                
                # حساب التكلفة يدوياً
                total_cost_val = 0
                for s in p.steps:
                    if s.employee and s.processing_time_minutes:
                        total_cost_val += (s.processing_time_minutes * s.employee.cost_per_minute)
                annual_cost_val = total_cost_val * p.annual_frequency
                
                # تقييم الحالة
                if flow_eff < 5:
                    status = "🔴 خطر"
                elif flow_eff < 20:
                    status = "🟠 سيء"
                elif flow_eff < 40:
                    status = "🟡 مقبول"
                else:
                    status = "🟢 جيد"

            summary_data.append({
                "العملية": p.name,
                "الفئة": p.category,
                "⏳ وقت الانتظار (دقيقة)": f"{wait:,.0f}",
                "⚡ كفاءة التدفق": f"{flow_eff:.1f}%",
                "🕐 زمن الدورة (ساعة)": f"{lead_time/60:.1f}",
                "💸 التكلفة السنوية (د.أ)": f"{annual_cost_val:,.2f}",
                "الحالة": status
            })
        
        df_summary = pd.DataFrame(summary_data)
        st.dataframe(df_summary, use_container_width=True)

        st.markdown("---")

        # --- 3. أهم 3 عمليات تحتاج تدخلاً (Pareto Mini) ---
        st.subheader("🚨 أهم 3 عمليات تحتاج تدخلاً فورياً")
        pareto_data = []
        for proc in all_processes:
            with app.app_context():
                p = Process.query.get(proc.id)
                wait = sum((s.wait_time_minutes or 0) for s in p.steps)
            pareto_data.append({"name": p.name, "waste": wait})
        
        df_pareto = pd.DataFrame(pareto_data).sort_values(by="waste", ascending=False).head(3)
        
        if not df_pareto.empty:
            cols = st.columns(3)
            for i, (_, row) in enumerate(df_pareto.iterrows()):
                with cols[i]:
                    st.error(f"**#{i+1}: {row['name']}**")
                    st.metric("⏳ وقت الهدر", f"{row['waste']:,.0f} دقيقة")
                    st.caption("ابدأ بتحسين هذه العملية فوراً.")

    else:
        st.info("لا توجد عمليات بعد. أضف عمليات من القائمة الجانبية لعرض لوحة القيادة.")

# ================== رحلة متلقي الخدمة ==================
elif menu == "رحلة متلقي الخدمة":
    st.subheader("خريطة رحلة متلقي الخدمة")
    st.markdown("هذه الخريطة توضح رحلة المعاملة من وجهة نظر المواطن او متلقي الخدمة")
    
    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية لعرض رحلتها", pnames)
        pid = int(sel.split(" - ")[0])
        steps = get_steps(pid)
        process = get_process_by_id(pid)
        
        if steps:
            step_labels = []
            proc_times = []
            wait_times = []
            step_types = []
            pain_scores = []
            
            for s in steps:
                with app.app_context():
                    emp = Employee.query.get(s.employee_id)
                    emp_title = emp.title if emp else "غير محدد"
                
                label = f"{s.step_order}. {s.step_name} - {emp_title}"
                step_labels.append(label)
                proc_times.append(s.processing_time_minutes or 0)
                wait_times.append(s.wait_time_minutes or 0)
                step_types.append(s.step_type)
                
                wait = s.wait_time_minutes or 0
                if wait > 2880:
                    pain_scores.append(10)
                elif wait > 1440:
                    pain_scores.append(8)
                elif wait > 480:
                    pain_scores.append(6)
                elif wait > 60:
                    pain_scores.append(4)
                elif wait > 0:
                    pain_scores.append(2)
                else:
                    pain_scores.append(0)
            
            st.subheader("مراحل الرحلة ونقاط الالم")
            
            fig_journey = go.Figure()
            fig_journey.add_trace(go.Bar(
                name='وقت العمل',
                y=step_labels,
                x=proc_times,
                orientation='h',
                marker=dict(color='#2ecc71'),
                text=[f"{p} دقيقة" if p > 0 else "" for p in proc_times],
                textposition='inside'
            ))
            fig_journey.add_trace(go.Bar(
                name='وقت الانتظار',
                y=step_labels,
                x=wait_times,
                orientation='h',
                marker=dict(color='#e74c3c'),
                text=[f"{w} دقيقة" if w > 0 else "" for w in wait_times],
                textposition='inside'
            ))
            fig_journey.update_layout(
                barmode='stack',
                height=400 + len(steps) * 50,
                title="رحلة المعاملة",
                xaxis_title="الوقت (دقيقة)",
                yaxis=dict(autorange="reversed")
            )
            st.plotly_chart(fig_journey, use_container_width=True)
            
            # نقاط الالم
            st.subheader("نقاط الالم في الرحلة")
            pain_steps = []
            for i, s in enumerate(steps):
                if pain_scores[i] >= 6:
                    with app.app_context():
                        emp = Employee.query.get(s.employee_id)
                        emp_title = emp.title if emp else "غير محدد"
                    pain_steps.append({
                        "الخطوة": s.step_name,
                        "المسؤول": emp_title,
                        "وقت الانتظار": f"{s.wait_time_minutes:.0f} دقيقة",
                        "درجة الالم": f"{pain_scores[i]}/10",
                        "فئة الهدر": s.waste_category or "غير محدد"
                    })
            
            if pain_steps:
                df_pain = pd.DataFrame(pain_steps)
                st.dataframe(df_pain, use_container_width=True)
                st.error(f"تم اكتشاف {len(pain_steps)} نقاط الم حرجة")
            else:
                st.success("لا توجد نقاط الم حرجة")
            
            # CSAT
            st.subheader("مؤشر رضا متلقي الخدمة التقديري")
            total_proc = sum(proc_times)
            total_wait = sum(wait_times)
            total_time = total_proc + total_wait
            
            if total_time > 0:
                wait_ratio = total_wait / total_time
                csat = max(0, 100 - (wait_ratio * 100))
            else:
                csat = 100
            
            col1, col2 = st.columns([1, 3])
            with col1:
                if csat >= 80:
                    st.markdown(f"<h1 style='color:green; text-align:center;'>{csat:.0f}%</h1>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align:center;'>ممتاز</p>", unsafe_allow_html=True)
                elif csat >= 50:
                    st.markdown(f"<h1 style='color:orange; text-align:center;'>{csat:.0f}%</h1>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align:center;'>متوسط</p>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<h1 style='color:red; text-align:center;'>{csat:.0f}%</h1>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align:center;'>ضعيف</p>", unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                **كيف يحسب هذا المؤشر؟**
                - مؤشر رضا متلقي الخدمة = 100% - (نسبة وقت الانتظار من اجمالي زمن الرحلة)
                - اجمالي وقت الرحلة: **{total_time:.0f} دقيقة**
                - وقت العمل الفعلي: **{total_proc:.0f} دقيقة**
                - وقت الانتظار: **{total_wait:.0f} دقيقة**
                """)
            
            # توصيات
            st.subheader("توصيات لتحسين رحلة متلقي الخدمة")
            recommendations = []
            for s in steps:
                with app.app_context():
                    emp = Employee.query.get(s.employee_id)
                    emp_title = emp.title if emp else "غير محدد"
                if s.step_type == 'NVA' and (s.wait_time_minutes or 0) > 1440:
                    recommendations.append(f"اتمتة خطوة '{s.step_name}' (المسؤول: {emp_title}): انتظار {s.wait_time_minutes:.0f} دقيقة يمكن اختزاله")
                elif s.step_type == 'NVA' and (s.wait_time_minutes or 0) > 0:
                    recommendations.append(f"تسريع خطوة '{s.step_name}' (المسؤول: {emp_title})")
                elif s.system_used == 'ورقي':
                    recommendations.append(f"رقمنة خطوة '{s.step_name}' (المسؤول: {emp_title})")
            
            if recommendations:
                for r in recommendations:
                    st.markdown(r)
            else:
                st.success("لا توجد توصيات حالية")
        else:
            st.info("لا توجد خطوات لهذه العملية")
    else:
        st.info("لا توجد عمليات بعد")

# ================== مصفوفة الاثر والتاثير ==================
elif menu == "مصفوفة الاثر والتاثير":
    st.subheader("مصفوفة الاثر والتاثير")
    st.markdown("هذه الاداة تساعدك على التخطيط لتعديل خطوة معينة وتوقع الاثر على باقي اجزاء المنظومة قبل التنفيذ")
    
    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel_process = st.selectbox("اختر العملية المستهدفة", pnames)
        pid = int(sel_process.split(" - ")[0])
        steps = get_steps(pid)
        
        if steps:
            step_names = [f"{s.step_order}. {s.step_name}" for s in steps]
            
            if "previous_step" not in st.session_state:
                st.session_state.previous_step = step_names[0] if step_names else None
            
            sel_step = st.selectbox("اختر الخطوة التي تريد تعديلها", step_names)
            
            if sel_step != st.session_state.previous_step:
                st.session_state.previous_step = sel_step
                st.session_state.pop("change_desc", None)
                st.rerun()
            
            st.markdown("---")
            st.markdown("### التعديل المقترح")
            change_description = st.text_area("صف التعديل الذي تخطط له", value=st.session_state.get("change_desc", ""))
            
            st.markdown("---")
            st.markdown("### الاطراف المتاثرة بالتعديل")
            
            impacted_parties = ["الموظف", "النظام (Oracle/GFMIS)", "جهة خارجية", "متلقي الخدمة"]
            impacts = {}
            
            for party in impacted_parties:
                with st.expander(f"{party}"):
                    col1, col2 = st.columns([1, 2])
                    with col1:
                        impact_type = st.selectbox("نوع الاثر", ["لا يوجد تغيير", "تغيير في الصلاحية", "تغيير في الوقت", "تغيير في التكلفة", "مقاومة متوقعة"], key=f"type_{party}")
                    with col2:
                        impact_desc = st.text_area("وصف الاثر", key=f"desc_{party}")
                    impacts[party] = {"type": impact_type, "desc": impact_desc}
            
            st.markdown("---")
            st.markdown("### تقدير مبدئي للعائد")
            selected_step_index = int(sel_step.split(".")[0]) - 1
            if selected_step_index < len(steps):
                target_step = steps[selected_step_index]
                current_wait = target_step.wait_time_minutes or 0
                if current_wait > 0:
                    saved_minutes = current_wait * 0.8
                    st.success(f"بتعديل هذه الخطوة (التي تنتظر {current_wait:.0f} دقيقة)، يمكنك توفير ما يقدر بـ **{saved_minutes:.0f} دقيقة** على الاقل لكل معاملة")
                else:
                    st.info("لا يوجد وقت انتظار مسجل لهذه الخطوة")
            
            if st.button("حفظ التحليل"):
                st.success("تم تسجيل التحليل بنجاح")
        else:
            st.info("لا توجد خطوات لهذه العملية")
    else:
        st.info("لا توجد عمليات بعد")

# ================== رفع ملف عمليات ==================
# ================== رفع ملف عمليات ==================
# ================== رفع ملف عمليات (تلقائي) ==================
elif menu == "رفع ملف عمليات":
    st.subheader("رفع ملف عمليات (Excel/CSV)")
    st.markdown("ارفع ملف Excel أو CSV. سيقرأ النظام أول ورقتين تلقائياً.")

    uploaded_file = st.file_uploader("اختر ملف Excel أو CSV", type=["xlsx", "csv"])

    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                df_process = pd.read_csv(uploaded_file, nrows=1)
                df_steps = pd.read_csv(uploaded_file, skiprows=2)
            else:
                # قراءة جميع أسماء الأوراق
                xl = pd.ExcelFile(uploaded_file)
                sheet_names = xl.sheet_names
                
                if len(sheet_names) >= 2:
                    df_process = pd.read_excel(uploaded_file, sheet_name=sheet_names[0])
                    df_steps = pd.read_excel(uploaded_file, sheet_name=sheet_names[1])
                else:
                    st.error("يجب أن يحتوي الملف على ورقتين على الأقل.")
                    st.stop()

            st.subheader("معاينة البيانات المستوردة")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**بيانات العملية:**")
                st.dataframe(df_process, use_container_width=True)
            with col2:
                st.markdown("**بيانات الخطوات:**")
                st.dataframe(df_steps, use_container_width=True)

            if st.button("تاكيد واستيراد البيانات", use_container_width=True):
                with app.app_context():
                    # استيراد العملية
                    process_row = df_process.iloc[0]
                    process_name = str(process_row['اسم العملية'])
                    category = str(process_row.get('الفئة', 'روتينية'))
                    freq = int(process_row.get('التكرار السنوي', 1))
                    status = str(process_row.get('الحالة', 'تحت_الدراسة'))

                    existing = Process.query.filter_by(name=process_name).first()
                    if existing:
                        process_id = existing.id
                        st.info(f"العملية '{process_name}' موجودة مسبقاً. ستُضاف إليها الخطوات.")
                    else:
                        new_process = Process(name=process_name, category=category,
                                             annual_frequency=freq, status=status)
                        db.session.add(new_process)
                        db.session.commit()
                        process_id = new_process.id

                    # استيراد الخطوات
                    steps_added = 0
                    for _, row in df_steps.iterrows():
                        step_order = int(row['رقم الترتيب'])
                        step_name = str(row['اسم الخطوة'])
                        proc_time = float(row.get('وقت المعالجة', 0))
                        wait_time = float(row.get('وقت الانتظار', 0))
                        step_type = str(row.get('النوع', 'VA'))
                        system_used = str(row.get('النظام', ''))
                        waste_cat = str(row.get('فئة الهدر', '')) if pd.notna(row.get('فئة الهدر', None)) else ''
                        emp_name = str(row.get('الموظف', ''))

                        # البحث عن الموظف أو إنشاؤه
                        employee = Employee.query.filter_by(title=emp_name).first()
                        if not employee and emp_name:
                            employee = Employee(title=emp_name, monthly_cost=500)
                            db.session.add(employee)
                            db.session.commit()

                        emp_id = employee.id if employee else None

                        new_step = Step(
                            process_id=process_id,
                            employee_id=emp_id,
                            step_order=step_order,
                            step_name=step_name,
                            processing_time_minutes=proc_time,
                            wait_time_minutes=wait_time,
                            step_type=step_type,
                            system_used=system_used,
                            waste_category=waste_cat
                        )
                        db.session.add(new_step)
                        steps_added += 1

                    db.session.commit()

                st.success(f"تم استيراد العملية '{process_name}' بنجاح مع {steps_added} خطوة!")
                st.balloons()

        except Exception as e:
            st.error("حدث خطأ أثناء قراءة الملف. تأكد من التنسيق الصحيح.")
            st.code(str(e))
# ================== تحليل باريتو ==================
elif menu == "🎯 تحليل باريتو (80/20)":
    st.subheader("🎯 تحليل باريتو للهدر في العمليات")
    st.markdown("هذا التحليل يطبق قاعدة 80/20 لمساعدتك على التركيز على العمليات الأكثر هدراً.")

    all_processes = get_processes()
    if all_processes:
        # 1. تجهيز البيانات
        pareto_data = []
        for proc in all_processes:
            with app.app_context():
                p = Process.query.get(proc.id)
                # حساب الهدر (وقت الانتظار) لكل عملية
                total_wait = sum((s.wait_time_minutes or 0) for s in p.steps)
            pareto_data.append({"name": p.name, "waste": total_wait})
        
        # 2. الترتيب التنازلي
        df_pareto = pd.DataFrame(pareto_data).sort_values(by="waste", ascending=False)
        df_pareto["نسبة مئوية"] = (df_pareto["waste"] / df_pareto["waste"].sum()) * 100
        df_pareto["تراكمي"] = df_pareto["نسبة مئوية"].cumsum()
        
        # 3. الرسم
        import plotly.graph_objects as go
        fig = go.Figure()
        fig.add_trace(go.Bar(name="وقت الهدر (دقيقة)", x=df_pareto["name"], y=df_pareto["waste"], marker_color="#e74c3c"))
        fig.add_trace(go.Scatter(name="النسبة التراكمية %", x=df_pareto["name"], y=df_pareto["تراكمي"], yaxis="y2", marker_color="#3498db"))
        fig.update_layout(
            title="مخطط باريتو",
            xaxis_title="العمليات",
            yaxis_title="وقت الهدر (دقيقة)",
            yaxis2=dict(title="النسبة التراكمية %", overlaying="y", side="right", range=[0, 100]),
            legend=dict(x=0.01, y=0.99)
        )
        st.plotly_chart(fig)

        # 4. التوصية التلقائية
        st.markdown("---")
        st.subheader("💡 توصية باريتو")
        # تحديد العمليات التي تشكل 80% من الهدر
        vital_few = df_pareto[df_pareto["تراكمي"] <= 80]
        if not vital_few.empty:
            names = "، ".join(vital_few["name"].tolist())
            st.success(f"🎯 **قاعدة 80/20:** التركيز على تحسين هذه العمليات ({names}) سيعالج ما يقارب 80% من إجمالي الهدر.")
        else:
            st.info("حتى الآن، لا توجد عملية واحدة تستحوذ على 80% من الهدر لوحدها.")
    else:
        st.info("لا توجد عمليات بعد لتحليلها.")
        # ================== SIPOC ==================
elif menu == "📋 SIPOC":
    st.subheader("📋 تحليل SIPOC")
    st.markdown("نظرة شاملة للعملية: الجهة المعنية، المدخلات، العملية، المخرجات، متلقي الخدمة.")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية", pnames)
        pid = int(sel.split(" - ")[0])
        steps = get_steps(pid)

        if steps:
            # اقتراح تلقائي للبيانات
            suppliers = "جميع الجهات الحكومية (وزارات، دوائر)"
            inputs_list = "طلب مكتمل، مستندات ثبوتية، موافقات سابقة"
            
            # العملية = أسماء الخطوات
            process_desc = " → ".join([s.step_name for s in steps])
            
            outputs_list = "معاملة منجزة، إشعار إلكتروني"
            customers = "متلقي الخدمة، جهة رقابية، ديوان المحاسبة"

            with st.form("sipoc_form"):
                st.subheader("✏️ عدل جدول SIPOC")
                col1, col2 = st.columns(2)
                with col1:
                    sup = st.text_area("**S - الجهة المعنية (Suppliers)**", value=suppliers, height=100)
                    inp = st.text_area("**I - المدخلات (Inputs)**", value=inputs_list, height=100)
                    proc = st.text_area("**P - العملية (Process)**", value=process_desc, height=100)
                with col2:
                    outp = st.text_area("**O - المخرجات (Outputs)**", value=outputs_list, height=100)
                    cust = st.text_area("**C - متلقي الخدمة (Customers)**", value=customers, height=100)

                if st.form_submit_button("💾 حفظ وعرض SIPOC"):
                    st.success("تم تحديث SIPOC")
                    
                    # عرض الجدول النهائي
                    df_sipoc = pd.DataFrame({
                        "المكون": ["S - الجهة المعنية", "I - المدخلات", "P - العملية", "O - المخرجات", "C - متلقي الخدمة"],
                        "الوصف": [sup, inp, proc, outp, cust]
                    })
                    st.dataframe(df_sipoc, use_container_width=True)
        else:
            st.info("لا توجد خطوات لهذه العملية.")
    else:
        st.info("لا توجد عمليات بعد.")
        # ================== RACI Matrix ==================
elif menu == "📊 RACI":
    st.subheader("📊 مصفوفة المسؤوليات (RACI)")
    st.markdown("تحديد المسؤوليات لكل خطوة: Responsible, Accountable, Consulted, Informed")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية", pnames)
        pid = int(sel.split(" - ")[0])
        steps = get_steps(pid)

        if steps:
            # جمع كل الموظفين المشاركين في العملية
            unique_emps = {}
            for s in steps:
                with app.app_context():
                    emp = Employee.query.get(s.employee_id)
                    if emp:
                        unique_emps[emp.id] = emp.title

            if unique_emps:
                # اقتراح تلقائي لـ RACI
                st.markdown("### اقتراح آلي (يمكنك تعديله)")
                
                raci_data = []
                for s in steps:
                    with app.app_context():
                        emp = Employee.query.get(s.employee_id)
                        emp_title = emp.title if emp else "-"
                    
                    row = {
                        "الخطوة": s.step_name,
                        "R (مسؤول)": emp_title,
                        "A (معتمد)": "مدير عام" if "توقيع" in s.step_name else "مدير مالي",
                        "C (مُستشار)": "مدقق مالي" if "تدقيق" in s.step_name else "-",
                        "I (مُبلّغ)": "جميع الأطراف"
                    }
                    raci_data.append(row)
                
                df_raci = pd.DataFrame(raci_data)
                
                # عرض قابل للتعديل
                with st.form("raci_form"):
                    edited_rows = []
                    for i, row in enumerate(raci_data):
                        st.markdown(f"**{row['الخطوة']}**")
                        col_r, col_a, col_c, col_i = st.columns(4)
                        with col_r:
                            new_r = st.text_input("R", value=row["R (مسؤول)"], key=f"r_{i}")
                        with col_a:
                            new_a = st.text_input("A", value=row["A (معتمد)"], key=f"a_{i}")
                        with col_c:
                            new_c = st.text_input("C", value=row["C (مُستشار)"], key=f"c_{i}")
                        with col_i:
                            new_i = st.text_input("I", value=row["I (مُبلّغ)"], key=f"i_{i}")
                        edited_rows.append({"الخطوة": row["الخطوة"], "R": new_r, "A": new_a, "C": new_c, "I": new_i})
                    
                    if st.form_submit_button("💾 حفظ RACI"):
                        st.success("تم حفظ مصفوفة RACI")
                        st.dataframe(pd.DataFrame(edited_rows), use_container_width=True)
            else:
                st.info("لا يوجد موظفين مرتبطين بهذه العملية.")
        else:
            st.info("لا توجد خطوات لهذه العملية.")
    else:
        st.info("لا توجد عمليات بعد.")
     # ================== الخريطة الحرارية ==================
elif menu == "🗺️ الخريطة الحرارية":
    st.subheader("🗺️ الخريطة الحرارية للعمليات")
    st.markdown("نظرة شاملة على صحة جميع العمليات. ركز على العمليات الحمراء أولاً.")

    all_processes = get_processes()
    if all_processes:
        heatmap_data = []
        for proc in all_processes:
            with app.app_context():
                p = Process.query.get(proc.id)
                total_wait = sum((s.wait_time_minutes or 0) for s in p.steps)
                total_processing = sum((s.processing_time_minutes or 0) for s in p.steps)
                lead_time = total_processing + total_wait
                flow_eff = (total_processing / lead_time * 100) if lead_time > 0 else 100
                annual_cost = p.annual_cost

            # تحديد مستوى الخطر
            if flow_eff < 5:
                risk = "🔴 خطر"
            elif flow_eff < 20:
                risk = "🟠 سيء"
            elif flow_eff < 40:
                risk = "🟡 مقبول"
            else:
                risk = "🟢 جيد"

            heatmap_data.append({
                "العملية": p.name,
                "الفئة": p.category,
                "كفاءة التدفق %": f"{flow_eff:.1f}%",
                "زمن الدورة (ساعة)": round(lead_time / 60, 1),
                "وقت الانتظار (ساعة)": round(total_wait / 60, 1),
                "التكلفة الشاملة (د.أ)": round(annual_cost + (total_wait * 0.1), 0),
                "مستوى الخطر": risk
            })
        
        df_heatmap = pd.DataFrame(heatmap_data)
        st.dataframe(df_heatmap, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 🎯 أولويات التحسين")
        st.markdown("- 🔴 **خطر:** ابدأ بهذه العمليات فوراً")
        st.markdown("- 🟠 **سيء:** خطط لتحسينها هذا الشهر")
        st.markdown("- 🟡 **مقبول:** جدولها للربع القادم")
        st.markdown("- 🟢 **جيد:** راقبها وحافظ على أدائها")
    else:
        st.info("لا توجد عمليات بعد.")
        # ================== توصيات التحسين ==================
elif menu == "🚀 توصيات التحسين":
    st.subheader("🚀 توصيات التحسين (To-Be)")
    st.markdown("خطط للحالة المستقبلية لكل عملية واحسب العائد المتوقع من التحسين.")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية", pnames)
        pid = int(sel.split(" - ")[0])
        process = get_process_by_id(pid)
        steps = get_steps(pid)

        if process and steps:
            with app.app_context():
                p = Process.query.get(pid)
                current_eff = p.flow_efficiency
                current_lead = p.lead_time_minutes
                current_cost = p.annual_cost
                
                total_wait = sum((s.wait_time_minutes or 0) for s in steps)
                total_proc = sum((s.processing_time_minutes or 0) for s in steps)

            st.markdown("---")
            st.markdown("### 📊 الحالة الحالية (As-Is)")
            col1, col2, col3 = st.columns(3)
            col1.metric("كفاءة التدفق الحالية", f"{current_eff:.1f}%")
            col2.metric("زمن الدورة الحالي", f"{current_lead/60:.1f} ساعة")
            col3.metric("التكلفة السنوية الحالية", f"{current_cost:,.2f} د.أ")

            st.markdown("---")
            st.markdown("### 🎯 الحالة المستقبلية المستهدفة (To-Be)")
            
            with st.form("tobe_form"):
                target_eff = st.slider("كفاءة التدفق المستهدفة (%)", min_value=float(current_eff), max_value=100.0, value=min(float(current_eff) * 5, 85.0), step=1.0)
                target_lead_hours = st.number_input("زمن الدورة المستهدف (ساعة)", min_value=0.1, value=max(0.5, float(current_lead/60) * 0.2), step=0.5)
                target_cost = st.number_input("التكلفة السنوية المستهدفة (د.أ)", min_value=0.0, value=float(current_cost) * 0.3, step=100.0)
                
                improvement_desc = st.text_area("وصف خطة التحسين", placeholder="مثلاً: تفعيل التوقيع الإلكتروني في Oracle، إلغاء الأرشفة الورقية، دمج خطوتين...")
                
                if st.form_submit_button("💾 حفظ التوصية"):
                    # حساب العائد
                    saving = current_cost - target_cost
                    roi = (saving / (current_cost + 1)) * 100
                    
                    st.success("تم حفظ التوصية!")
                    
                    st.markdown("---")
                    st.markdown("### 📈 ملخص العائد المتوقع")
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("الوفر السنوي المتوقع", f"{saving:,.2f} د.أ")
                    c2.metric("تحسين كفاءة التدفق", f"+{target_eff - current_eff:.1f}%")
                    c3.metric("تقليص زمن الدورة", f"-{current_lead/60 - target_lead_hours:.1f} ساعة")
                    
                    st.markdown("---")
                    st.markdown("### 🖨️ تقرير جاهز للطباعة")
                    st.info(f"""
                    **تقرير تحسين عملية: {process.name}**
                    
                    **الوضع الحالي:**
                    - كفاءة التدفق: {current_eff:.1f}%
                    - زمن الدورة: {current_lead/60:.1f} ساعة
                    - التكلفة السنوية: {current_cost:,.2f} د.أ
                    
                    **الوضع المستهدف:**
                    - كفاءة التدفق: {target_eff:.1f}%
                    - زمن الدورة: {target_lead_hours:.1f} ساعة
                    - التكلفة السنوية: {target_cost:,.2f} د.أ
                    
                    **خطة التحسين:** {improvement_desc}
                    
                    **العائد المتوقع:**
                    - وفر سنوي: {saving:,.2f} د.أ
                    - تحسين الكفاءة: +{target_eff - current_eff:.1f}%
                    """)
        else:
            st.info("لا توجد خطوات لهذه العملية.")
    else:
        st.info("لا توجد عمليات بعد.")
# ================== مخطط BPMN (مبسط وواضح) ==================
elif menu == "📊 مخطط BPMN":
    st.subheader("📊 مخطط مسار العمل (BPMN)")
    st.markdown("رسم تخطيطي مبسط لسير العملية.")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية لرسم مخططها", pnames)
        pid = int(sel.split(" - ")[0])
        steps = get_steps(pid)
        process = get_process_by_id(pid)

        if process and steps:
            # عرض مبسط كجدول ملون بدلاً من رسم معقد
            st.markdown("---")
            st.markdown(f"### 📋 مخطط: {process.name}")
            
            for i, s in enumerate(steps):
                with app.app_context():
                    emp = Employee.query.get(s.employee_id)
                    emp_title = emp.title if emp else "-"
                
                # تحديد اللون
                if s.step_type == 'VA':
                    color = '#2ecc71'
                    emoji = '✅'
                elif s.step_type == 'BNVA':
                    color = '#f39c12'
                    emoji = '⚠️'
                else:
                    color = '#e74c3c'
                    emoji = '❌'
                
                # عرض الخطوة كبطاقة ملونة
                st.markdown(f"""
                <div style="
                    border-left: 5px solid {color};
                    background-color: {color}15;
                    padding: 10px;
                    margin: 5px 0;
                    border-radius: 5px;
                ">
                    <b>{emoji} {s.step_order}. {s.step_name}</b> ({s.step_type})<br>
                    <small>👤 {emp_title} | ⏱️ عمل: {s.processing_time_minutes}د | انتظار: {s.wait_time_minutes}د</small>
                </div>
                """, unsafe_allow_html=True)
                
                # سهم بين الخطوات
                if i < len(steps) - 1:
                    st.markdown("<div style='text-align:center; color:#7f8c8d;'>⬇️</div>", unsafe_allow_html=True)
            
                                # وسيلة الإيضاح
            st.markdown("---")
            st.markdown("### 🔍 وسيلة الإيضاح")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown("""
                <div style="background-color:#2ecc71; padding:10px; border-radius:5px; color:white; text-align:center;">
                    <b>VA (قيمة مضافة)</b><br>
                    <small>متلقي الخدمة يدفع مقابلها</small>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown("""
                <div style="background-color:#f39c12; padding:10px; border-radius:5px; color:black; text-align:center;">
                    <b>BNVA (ضرورية)</b><br>
                    <small>إجراء قانوني / رقابي</small>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown("""
                <div style="background-color:#e74c3c; padding:10px; border-radius:5px; color:white; text-align:center;">
                    <b>NVA (هدر)</b><br>
                    <small>يمكن ويجب إلغاؤها</small>
                </div>
                """, unsafe_allow_html=True)
    else:
        st.info("لا توجد عمليات بعد.")
        # ================== تقرير العملية الشامل ==================
elif menu == "📄 تقرير العملية":
    st.subheader("📄 تقرير العملية الشامل")
    st.markdown("بطاقة تقرير كاملة للعملية المختارة، جاهزة للطباعة أو الحفظ.")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية لإنشاء تقريرها", pnames)
        pid = int(sel.split(" - ")[0])
        process = get_process_by_id(pid)
        steps = get_steps(pid)

        if process and steps:
            with app.app_context():
                p = Process.query.get(pid)
                wait = sum((s.wait_time_minutes or 0) for s in p.steps)
                proc_time = sum((s.processing_time_minutes or 0) for s in p.steps)
                lead_time = proc_time + wait
                flow_eff = (proc_time / lead_time * 100) if lead_time > 0 else 100
                
                # حساب التكلفة
                total_cost_val = 0
                for s in p.steps:
                    if s.employee and s.processing_time_minutes:
                        total_cost_val += (s.processing_time_minutes * s.employee.cost_per_minute)
                annual_cost_val = total_cost_val * p.annual_frequency
                
                # SIPOC تلقائي
                suppliers = "جميع الجهات الحكومية"
                inputs_list = "طلب مكتمل، مستندات ثبوتية"
                process_desc = " → ".join([s.step_name for s in steps])
                outputs_list = "معاملة منجزة، إشعار"
                customers = "متلقي الخدمة، جهة رقابية"

            # --- تجميع التقرير في بطاقة واحدة ---
            with st.expander(f"📄 تقرير: {process.name}", expanded=True):
                st.markdown(f"## 🏷️ بطاقة تعريف العملية (SIPOC)")
                col_s, col_i = st.columns(2)
                with col_s:
                    st.markdown(f"**S - الجهة المعنية:** {suppliers}")
                    st.markdown(f"**I - المدخلات:** {inputs_list}")
                    st.markdown(f"**P - العملية:** {process_desc}")
                with col_i:
                    st.markdown(f"**O - المخرجات:** {outputs_list}")
                    st.markdown(f"**C - متلقي الخدمة:** {customers}")
                
                st.markdown("---")
                st.markdown(f"## 📊 مؤشرات الأداء الرئيسية")
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("كفاءة التدفق", f"{flow_eff:.1f}%")
                col2.metric("زمن الدورة", f"{lead_time/60:.1f} ساعة")
                col3.metric("وقت الانتظار", f"{wait/60:.1f} ساعة")
                col4.metric("التكلفة السنوية", f"{annual_cost_val:,.2f} د.أ")
                
                st.markdown("---")
                st.markdown(f"## 📋 خطوات العملية ({len(steps)} خطوة)")
                steps_data = []
                for s in steps:
                    with app.app_context():
                        emp = Employee.query.get(s.employee_id)
                        emp_title = emp.title if emp else "-"
                    steps_data.append({
                        "#": s.step_order,
                        "الخطوة": s.step_name,
                        "النوع": s.step_type,
                        "الموظف": emp_title,
                        "وقت العمل": f"{s.processing_time_minutes}د",
                        "وقت الانتظار": f"{s.wait_time_minutes}د",
                        "النظام": s.system_used or "-"
                    })
                st.dataframe(pd.DataFrame(steps_data), use_container_width=True)
                
                st.markdown("---")
                st.markdown(f"## 🗺️ مخطط التدفق")
                for i, s in enumerate(steps):
                    with app.app_context():
                        emp = Employee.query.get(s.employee_id)
                        emp_title = emp.title if emp else "-"
                    
                    if s.step_type == 'VA':
                        color = '#2ecc71'
                        emoji = '✅'
                    elif s.step_type == 'BNVA':
                        color = '#f39c12'
                        emoji = '⚠️'
                    else:
                        color = '#e74c3c'
                        emoji = '❌'
                    
                    st.markdown(f"""
                    <div style="border-left:5px solid {color}; background-color:{color}15; padding:8px; margin:3px 0; border-radius:5px;">
                        <b>{emoji} {s.step_order}. {s.step_name}</b> ({s.step_type})<br>
                        <small>👤 {emp_title} | ⏱️ عمل: {s.processing_time_minutes}د | انتظار: {s.wait_time_minutes}د</small>
                    </div>
                    """, unsafe_allow_html=True)
                    if i < len(steps) - 1:
                        st.markdown("<div style='text-align:center;'>⬇️</div>", unsafe_allow_html=True)
                
                st.markdown("---")
                st.markdown(f"## 💡 توصيات آلية للتحسين")
                recommendations = []
                for s in steps:
                    if s.step_type == 'NVA' and (s.wait_time_minutes or 0) > 1440:
                        recommendations.append(f"🚀 **أتمتة '{s.step_name}'**: توفير {s.wait_time_minutes} دقيقة عبر التوقيع الإلكتروني.")
                    elif s.system_used == 'ورقي':
                        recommendations.append(f"📄 **رقمنة '{s.step_name}'**: تحويلها إلى إلكترونية.")
                if recommendations:
                    for r in recommendations:
                        st.markdown(r)
                else:
                    st.success("لا توجد توصيات حرجة.")
                
                # زر تحميل التقرير (كنص)
                report_text = f"""
                تقرير العملية: {process.name}
                ================================
                كفاءة التدفق: {flow_eff:.1f}%
                زمن الدورة: {lead_time/60:.1f} ساعة
                وقت الانتظار: {wait/60:.1f} ساعة
                التكلفة السنوية: {annual_cost_val:,.2f} د.أ
                
                خطوات العملية:
                {chr(10).join([f"{s.step_order}. {s.step_name} ({s.step_type})" for s in steps])}
                
                التوصيات:
                {chr(10).join(recommendations) if recommendations else "لا توجد توصيات حرجة."}
                """
                st.download_button("📥 تحميل التقرير (نص)", data=report_text, file_name=f"تقرير_{process.name}.txt", mime="text/plain")
        else:
            st.info("لا توجد خطوات لهذه العملية.")
    else:
        st.info("لا توجد عمليات بعد.")
# ================== تقرير PDF (مصحح بالكامل) ==================
elif menu == "📄 تقرير PDF":
    st.subheader("📄 تصدير تقرير احترافي")
    st.markdown("اختر العملية وسيتم إنشاء تقرير HTML جاهز للتحميل والطباعة.")

    processes = get_processes()
    if processes:
        pnames = [f"{p.id} - {p.name}" for p in processes]
        sel = st.selectbox("اختر العملية لإنشاء تقريرها", pnames)
        pid = int(sel.split(" - ")[0])
        process = get_process_by_id(pid)
        steps = get_steps(pid)

        if process and steps:
            with app.app_context():
                p = Process.query.get(pid)
                wait = sum((s.wait_time_minutes or 0) for s in p.steps)
                proc_time = sum((s.processing_time_minutes or 0) for s in p.steps)
                lead_time = proc_time + wait
                flow_eff = (proc_time / lead_time * 100) if lead_time > 0 else 100

                # حساب التكلفة
                total_cost_val = 0
                for s in p.steps:
                    if s.employee and s.processing_time_minutes:
                        total_cost_val += (s.processing_time_minutes * s.employee.cost_per_minute)
                annual_cost_val = total_cost_val * p.annual_frequency

                # تجهيز بيانات الموظفين مسبقاً
                employee_names = {}
                for s in p.steps:
                    emp = Employee.query.get(s.employee_id)
                    employee_names[s.id] = emp.title if emp else "-"

            # تجهيز التوصيات
            rec_list = []
            for s in steps:
                if s.step_type == 'NVA' and (s.wait_time_minutes or 0) > 1440:
                    rec_list.append(f"<li>🚀 أتمتة '{s.step_name}': توفير {s.wait_time_minutes} دقيقة عبر التوقيع الإلكتروني.</li>")
                elif s.system_used == 'ورقي':
                    rec_list.append(f"<li>📄 رقمنة '{s.step_name}': تحويلها إلى إلكترونية.</li>")
            rec_html = "".join(rec_list) if rec_list else "<li>لا توجد توصيات حرجة.</li>"

            # تجهيز صفوف الجدول
            steps_rows = ""
            for i, s in enumerate(steps):
                emp_title = employee_names.get(s.id, "-")
                if s.step_type == 'NVA':
                    row_color = "#f8d7da"
                elif s.step_type == 'BNVA':
                    row_color = "#fff3cd"
                else:
                    row_color = "#d4edda"
                steps_rows += f"""
                <tr style="background-color: {row_color};">
                    <td>{i+1}</td>
                    <td>{s.step_name}</td>
                    <td>{s.step_type}</td>
                    <td>{emp_title}</td>
                    <td>{s.processing_time_minutes} دقيقة</td>
                    <td>{s.wait_time_minutes} دقيقة</td>
                </tr>"""

            # بناء التقرير
            html_report = f"""<!DOCTYPE html>
<html dir="rtl">
<head>
    <meta charset="UTF-8">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;700&display=swap');
        body {{ font-family: 'Tajawal', sans-serif; margin: 30px; color: #1e293b; }}
        .header {{ text-align: center; border-bottom: 3px solid #2563eb; padding-bottom: 15px; margin-bottom: 25px; }}
        .header h1 {{ color: #2563eb; margin: 0; font-size: 24px; }}
        .header p {{ color: #64748b; margin: 5px 0 0 0; }}
        .section-title {{ color: #2563eb; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px; margin-top: 25px; }}
        .kpi-box {{ display: inline-block; width: 22%; margin: 1%; padding: 15px; background-color: #f8fafc; border-radius: 8px; text-align: center; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .kpi-box h3 {{ margin: 0; font-size: 13px; color: #64748b; }}
        .kpi-box p {{ margin: 8px 0 0 0; font-size: 22px; font-weight: bold; color: #1e293b; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th {{ background-color: #1e293b; color: white; padding: 10px; }}
        td {{ padding: 8px; border: 1px solid #e2e8f0; text-align: center; }}
        .footer {{ margin-top: 30px; text-align: center; font-size: 12px; color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 10px; }}
        @media print {{ body {{ margin: 0; }} }}
    </style>
</head>
<body>
    <div class="header">
        <h1>تقرير تحليل عملية: {process.name}</h1>
        <p>للدوائر الحكومية- نظام إعادة هندسة العمليات</p>
    </div>

    <h2 class="section-title">📊 مؤشرات الأداء الرئيسية</h2>
    <div>
        <div class="kpi-box"><h3>كفاءة التدفق</h3><p>{flow_eff:.1f}%</p></div>
        <div class="kpi-box"><h3>زمن الدورة</h3><p>{lead_time/60:.1f} ساعة</p></div>
        <div class="kpi-box"><h3>وقت الانتظار</h3><p>{wait/60:.1f} ساعة</p></div>
        <div class="kpi-box"><h3>التكلفة السنوية</h3><p>{annual_cost_val:,.2f} د.أ</p></div>
    </div>

    <h2 class="section-title">📋 تفاصيل العملية (SIPOC)</h2>
    <p><b>S - الجهة المعنية:</b> جميع الجهات الحكومية</p>
    <p><b>I - المدخلات:</b> طلب مكتمل، مستندات ثبوتية</p>
    <p><b>P - العملية:</b> {' → '.join([s.step_name for s in steps])}</p>
    <p><b>O - المخرجات:</b> معاملة منجزة، إشعار</p>
    <p><b>C - متلقي الخدمة:</b> متلقي الخدمة، جهة رقابية</p>

    <h2 class="section-title">🗺️ خطوات العملية ({len(steps)} خطوة)</h2>
    <table>
        <tr><th>#</th><th>الخطوة</th><th>النوع</th><th>الموظف</th><th>وقت العمل</th><th>وقت الانتظار</th></tr>
        {steps_rows}
    </table>

    <h2 class="section-title">💡 توصيات التحسين</h2>
    <ul>{rec_html}</ul>

    <div class="footer">
        <p>تم إنشاء هذا التقرير بواسطة نظام إعادة هندسة العمليات © 2024</p>
    </div>
</body>
</html>"""

            # عرض معاينة مصغرة
            st.markdown("### 📄 معاينة التقرير")
            st.components.v1.html(html_report, height=500, scrolling=True)

            # زر تحميل التقرير
            st.download_button(
                label="📥 تحميل التقرير (HTML - افتحه واطبعه)",
                data=html_report,
                file_name=f"تقرير_{process.name}.html",
                mime="text/html",
                use_container_width=True
            )
            
            st.info("💡 بعد تحميل الملف، افتحه في المتصفح واضغط Ctrl+P للطباعة أو الحفظ كـ PDF.")

        else:
            st.info("لا توجد خطوات لهذه العملية.")
    else:
        st.info("لا توجد عمليات بعد.")
        # ================== رفع نموذج Word ==================
elif menu == "📄 رفع نموذج Word":
    st.subheader("📄 استيراد عملية من نموذج Word")
    st.markdown("ارفع ملف Word (نموذج توثيق العملية) وسيتم استخراج البيانات تلقائياً.")

    uploaded_docx = st.file_uploader("اختر ملف Word", type=["docx"])

    if uploaded_docx is not None:
        try:
            from docx import Document
            
            doc = Document(uploaded_docx)
            
            # استخراج النص من جميع الجداول
            extracted_data = {}
            
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        key = cells[0].text.strip()
                        value = cells[1].text.strip()
                        if key and value:
                            extracted_data[key] = value
            
            # استخراج النص من الفقرات (للبحث عن المدخلات والمخرجات)
            all_text = "\n".join([p.text for p in doc.paragraphs])
            
            st.markdown("---")
            st.markdown("### 📋 البيانات المستخرجة من الملف")
            
            # محاولة تخمين البيانات
            process_name = ""
            category = "روتينية"
            frequency = 1
            inputs_text = ""
            outputs_text = ""
            steps_text = ""
            
            for key, value in extracted_data.items():
                if "اسم العملية" in key:
                    process_name = value
                elif "تصنيف" in key or "فئة" in key:
                    category = value
                elif "تكرار" in key:
                    try:
                        frequency = int(value)
                    except:
                        frequency = 1
            
            # استخراج المدخلات والمخرجات من النص العام
            if "مدخلات" in all_text:
                input_start = all_text.find("مدخلات")
                input_end = all_text.find("مخرجات", input_start) if "مخرجات" in all_text else len(all_text)
                inputs_text = all_text[input_start:input_end].strip()
            
            if "مخرجات" in all_text:
                output_start = all_text.find("مخرجات")
                output_end = all_text.find("العمليات المرتبطة", output_start) if "العمليات المرتبطة" in all_text else len(all_text)
                outputs_text = all_text[output_start:output_end].strip()
            
            # عرض البيانات المستخرجة في نموذج قابل للتعديل
            with st.form("import_docx_form"):
                st.markdown("#### ✏️ عدل البيانات قبل الاستيراد")
                
                final_name = st.text_input("اسم العملية", value=process_name)
                col1, col2 = st.columns(2)
                with col1:
                    final_category = st.selectbox("الفئة", ["استراتيجية", "انتصار_سريع", "روتينية", "للدراسة"])
                    final_freq = st.number_input("التكرار السنوي", min_value=1, value=frequency)
                with col2:
                    final_status = st.selectbox("الحالة", ["غير_مبدوء", "تحت_الدراسة", "مكتمل"])
                
                st.text_area("المدخلات (تم استخراجها)", value=inputs_text, height=100)
                st.text_area("المخرجات (تم استخراجها)", value=outputs_text, height=100)
                
                # عرض الجدوال المستخرجة
                if extracted_data:
                    st.markdown("**البيانات المستخرجة من الجداول:**")
                    df_extracted = pd.DataFrame(list(extracted_data.items()), columns=["الحقل", "القيمة"])
                    st.dataframe(df_extracted, use_container_width=True)
                
                if st.form_submit_button("💾 استيراد العملية"):
                    if final_name:
                        add_process_to_db(final_name, final_category, final_freq, final_status)
                        st.success(f"تم استيراد العملية: {final_name}")
                        st.info("يمكنك الآن إضافة الخطوات يدوياً من صفحة 'إضافة خطوات'.")
                        st.rerun()
                    else:
                        st.error("الرجاء إدخال اسم العملية")
            
            # عرض النص الكامل للملف
            with st.expander("📝 عرض النص الكامل للملف"):
                st.text_area("النص المستخرج", all_text, height=300)
                
        except ImportError:
            st.error("⚠️ مكتبة python-docx غير مثبتة. أضفها إلى requirements.txt.")
        except Exception as e:
            st.error(f"حدث خطأ أثناء قراءة الملف: {e}")
# ================== دليل الاستخدام ==================
elif menu == "دليل الاستخدام":
    st.subheader("دليل استخدام نظام اعادة هندسة العمليات")
    
    with st.expander("1. الرئيسية", expanded=True):
        st.markdown("""
        **ماذا ترى؟** قائمة بجميع العمليات في صناديق قابلة للطي.
        - كل صندوق يعرض: كفاءة التدفق، التكلفة السنوية، التكرار السنوي.
        - **تعديل:** لتغيير اسم العملية، الفئة، التكرار، او الحالة.
        - **حذف:** لحذف العملية نهائيا.
        """)
    
    with st.expander("2. ادارة الموظفين"):
        st.markdown("""
        **اضافة موظف جديد:**
        - اكتب المسمى الوظيفي.
        - اختر نطاق الراتب من شريط التمرير.
        - يحسب التطبيق المتوسط تلقائيا ويستخدمه في حسابات التكلفة.
        
        **تعديل/حذف موظف:**
        - اضغط تعديل لتغيير المسمى او الراتب.
        - اضغط حذف لحذف الموظف.
        """)
    
    with st.expander("3. اضافة عملية"):
        st.markdown("""
        | الحقل | الوصف |
        |-------|-------|
        | اسم العملية | مثل: المناقلات المالية |
        | الفئة | استراتيجية، انتصار_سريع، روتينية، للدراسة |
        | التكرار السنوي | كم مرة تحدث في السنة |
        | الحالة | غير_مبدوء، تحت_الدراسة، مكتمل |
        """)
    
    with st.expander("4. اضافة خطوات"):
        st.markdown("""
        | الحقل | الوصف | مثال |
        |-------|-------|------|
        | العملية | اختر العملية المستهدفة | المناقلات المالية |
        | الموظف | من يقوم بهذه الخطوة | مدقق مالي |
        | رقم الترتيب | تسلسل الخطوة | 1, 2, 3... |
        | اسم الخطوة | وصف مختصر | انتظار توقيع المدير |
        | وقت المعالجة | دقائق العمل الفعلي | 5 |
        | وقت الانتظار | دقائق الانتظار | 1440 (يوم كامل) |
        | نوع الخطوة | VA = قيمة، BNVA = ضرورية، NVA = هدر |
        | النظام | Oracle, GFMIS, ورقي, Outlook... |
        | فئة الهدر | للخطوات NVA فقط | انتظار، موافقات_زائدة |
        """)
    
    with st.expander("5. لوحة التحكم"):
        st.markdown("""
        اختر عملية لتشاهد:
        - **3 بطاقات:** كفاءة التدفق، زمن الدورة، التكلفة السنوية.
        - **الرسم الشريطي:** اخضر = وقت عمل، احمر = وقت انتظار.
        - **الرسم الدائري:** العمل المفيد مقابل الهدر.
        - **جدول تفصيلي** لجميع الخطوات.
        """)
    
    with st.expander("6. كيف تقرا النتائج؟"):
        st.markdown("""
        | كفاءة التدفق | التقييم | الاجراء |
        |-------------|--------|---------|
        | اقل من 5% | كارثة | تدخل فوري |
        | 5% - 20% | سيء جدا | اولوية للتحسين |
        | 20% - 40% | مقبول | مجال للتحسين |
        | 40% - 60% | جيد | تحسينات طفيفة |
        | اكثر من 60% | ممتاز | حافظ على المستوى |
        """)
    with st.expander("💸 فهم التكاليف (العمل فقط مقابل الشاملة)"):
        st.markdown("""
        ### 📊 نوعان من التكلفة في لوحة التحكم
        
        **1. التكلفة السنوية (عمل فقط):**
        - تحتسب **وقت العمل الفعلي فقط** (Processing Time).
        - المعادلة: مجموع (وقت العمل لكل خطوة × تكلفة الموظف في الدقيقة) × التكرار السنوي.
        - هذا الرقم متحفظ ويعكس فقط الوقت الذي يعمل فيه الموظف فعلياً على المعاملة.
        
        **2. 💸 التكلفة الشاملة:**
        - تحتسب **وقت العمل + وقت الانتظار**.
        - المعادلة: (تكلفة العمل للتنفيذ الواحد + تكلفة الانتظار للتنفيذ الواحد) × التكرار السنوي.
        - هذا الرقم يعكس التكلفة الحقيقية للهدر، حيث أن الموظف يتقاضى راتبه حتى أثناء انتظار المعاملة.
        
        ### 🎯 مثال من عملية المناقلات المالية:
        | المؤشر | القيمة |
        |--------|--------|
        | وقت العمل الفعلي | 32 دقيقة |
        | وقت الانتظار | 4,320 دقيقة (3 أيام) |
        | التكلفة (عمل فقط) | ~686 د.أ سنوياً |
        | التكلفة الشاملة | ~103,000 د.أ سنوياً |
        
        ### 💡 للعرض على الإدارة:
        استخدم **التكلفة الشاملة** لإظهار الحجم الحقيقي للمشكلة.
        استخدم **كفاءة التدفق** لإظهار نسبة الهدر في الوقت.
        """)
    with st.expander("7. نصائح سريعة"):
        st.markdown("""
        - 1440 دقيقة = يوم عمل، 2880 = يومان، 10080 = اسبوع.
        - فئة **انتصار_سريع** = عمليات سهلة تبني سمعتك اولا.
        - كل عملية مكتملة = قصة نجاح للادارة.
        - الهدف: كشف الهدر ثم القضاء عليه.
        """)
